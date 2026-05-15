"""
app/document/assembler.py
──────────────────────────
Assembles a styled Word (.docx) sustainability report from:
  - GeneratedReport  : LLM-generated section text
  - ValidationReport : hallucination audit trail
  - SectionPlan      : company metadata and KPI data
"""

from __future__ import annotations

import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.core.report_template import STYLE, TITLE_PAGE, FOOTER, DOCUMENT_SECTION_ORDER
from app.core.config import settings



def _rgb(hex_str: str) -> RGBColor:
    """Convert 6-char hex string to RGBColor."""
    r = int(hex_str[0:2], 16)
    g = int(hex_str[2:4], 16)
    b = int(hex_str[4:6], 16)
    return RGBColor(r, g, b)



def _set_margins(doc: Document) -> None:
    """Apply margins from STYLE to all sections of the document."""
    for section in doc.sections:
        section.top_margin    = Cm(STYLE.margin_top)
        section.bottom_margin = Cm(STYLE.margin_bottom)
        section.left_margin   = Cm(STYLE.margin_left)
        section.right_margin  = Cm(STYLE.margin_right)



def _add_footer(doc: Document, company_name: str) -> None:
    """Add a simple text footer to the first document section."""
    section = doc.sections[0]
    footer = section.footer
    para = footer.paragraphs[0]
    para.clear()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(
        f"{FOOTER.left_text}  |  {company_name}"
    )
    run.font.size = Pt(FOOTER.font_size)
    run.font.color.rgb = _rgb(FOOTER.font_color)



def _add_title_page(
    doc: Document,
    company_name: str,
    reporting_year: str,
    industry: str,
) -> None:
    """Add a formatted title page."""

    # Top spacer
    for _ in range(6):
        doc.add_paragraph()

    # Main title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(TITLE_PAGE.main_title)
    title_run.font.name = STYLE.heading_font
    title_run.font.size = Pt(STYLE.title_size)
    title_run.font.color.rgb = _rgb(STYLE.heading_color)
    title_run.bold = True

    # Subtitle — company | year
    sub_text = TITLE_PAGE.subtitle_template.format(
        company_name=company_name,
        reporting_year=reporting_year,
    )
    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run(sub_text)
    sub_run.font.name = STYLE.heading_font
    sub_run.font.size = Pt(STYLE.heading1_size)
    sub_run.font.color.rgb = _rgb(STYLE.accent_color)

    # Industry
    ind_para = doc.add_paragraph()
    ind_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ind_run = ind_para.add_run(industry)
    ind_run.font.name = STYLE.body_font
    ind_run.font.size = Pt(STYLE.body_size)
    ind_run.font.color.rgb = _rgb("808080")

    # Spacer
    for _ in range(4):
        doc.add_paragraph()

    # Prepared by
    prep_para = doc.add_paragraph()
    prep_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    prep_run = prep_para.add_run(
        f"{TITLE_PAGE.prepared_by_label}: {TITLE_PAGE.prepared_by_value}"
    )
    prep_run.font.size = Pt(STYLE.caption_size)
    prep_run.font.color.rgb = _rgb("808080")

    # Date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(
        f"Generated: {date.today().strftime('%d %B %Y')}"
    )
    date_run.font.size = Pt(STYLE.caption_size)
    date_run.font.color.rgb = _rgb("808080")

    # Confidentiality notice
    for _ in range(2):
        doc.add_paragraph()
    conf_para = doc.add_paragraph()
    conf_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    conf_run = conf_para.add_run(TITLE_PAGE.confidentiality_notice)
    conf_run.font.size = Pt(STYLE.caption_size)
    conf_run.font.color.rgb = _rgb("C00000")
    conf_run.bold = True

    doc.add_page_break()



def _add_about_page(doc: Document, company_name: str, reporting_year: str) -> None:
    """Add a short 'About This Report' page."""
    h = doc.add_heading("About This Report", level=1)
    h.runs[0].font.color.rgb = _rgb(STYLE.heading_color)

    about_text = (
        f"This sustainability report for {company_name} covers the reporting "
        f"period ending {reporting_year}. It has been prepared in alignment with "
        f"the Global Reporting Initiative (GRI) Standards and the Task Force on "
        f"Climate-related Financial Disclosures (TCFD) recommendations.\n\n"
        f"Report content was generated using a large language model (LLM) pipeline "
        f"and automatically validated against source data to minimise numerical "
        f"inaccuracies. A validation summary is included as an appendix. All figures "
        f"should be independently verified before external publication."
    )
    para = doc.add_paragraph(about_text)
    para.runs[0].font.size = Pt(STYLE.body_size)
    doc.add_page_break()



def _add_section(
    doc: Document,
    section_id: str,
    title: str,
    gri_code: str,
    tcfd_pillar: str | None,
    body_text: str,
) -> None:
    """Add one GRI section with heading, metadata tags, and body text."""

    # Section heading
    h = doc.add_heading(title, level=1)
    if h.runs:
        h.runs[0].font.color.rgb = _rgb(STYLE.heading_color)
        h.runs[0].font.size = Pt(STYLE.heading1_size)

    # GRI / TCFD tag line
    tag_parts = [f"GRI: {gri_code}"]
    if tcfd_pillar:
        tag_parts.append(f"TCFD: {tcfd_pillar}")
    tag_para = doc.add_paragraph(" | ".join(tag_parts))
    tag_para.runs[0].font.size = Pt(STYLE.caption_size)
    tag_para.runs[0].font.color.rgb = _rgb(STYLE.accent_color)
    tag_para.runs[0].italic = True

    # Body text — split on double newlines into separate paragraphs
    paragraphs = [p.strip() for p in body_text.split("\n\n") if p.strip()]
    if not paragraphs:
        paragraphs = [body_text.strip()]

    for para_text in paragraphs:
        p = doc.add_paragraph(para_text)
        p.paragraph_format.space_after = Pt(STYLE.paragraph_space_after)
        if p.runs:
            p.runs[0].font.size = Pt(STYLE.body_size)
            p.runs[0].font.name = STYLE.body_font

    doc.add_paragraph()  



def _add_validation_appendix(doc: Document, validation_report) -> None:
    """Add a validation summary table as the final appendix."""
    doc.add_page_break()

    h = doc.add_heading("Appendix: Validation Summary", level=1)
    if h.runs:
        h.runs[0].font.color.rgb = _rgb(STYLE.heading_color)

    intro = doc.add_paragraph(
        f"The table below summarises the automated validation check performed "
        f"on all LLM-generated content. Each numeric value in the report was "
        f"cross-checked against the source CSV data within a "
        f"{int(settings.report_disclaimer and 5)}% tolerance threshold. "
        f"Flagged values require human review before publication."
    )
    if intro.runs:
        intro.runs[0].font.size = Pt(STYLE.body_size)

    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    headers = ["Section", "Numbers Checked", "Flags Raised", "Status"]
    for i, text in enumerate(headers):
        hdr_cells[i].text = text
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(STYLE.caption_size)
        run.font.color.rgb = _rgb("FFFFFF")

        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), STYLE.table_header_bg)
        shd.set(qn("w:color"), "FFFFFF")
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)

    for sr in validation_report.section_results:
        row_cells = table.add_row().cells
        status = "PASS ✓" if sr.passed else "REVIEW ⚠"
        row_data = [
            sr.section_title,
            str(len(sr.extracted_numbers)),
            str(sr.flag_count),
            status,
        ]
        for i, text in enumerate(row_data):
            row_cells[i].text = text
            run = row_cells[i].paragraphs[0].runs[0]
            run.font.size = Pt(STYLE.caption_size)
            if i == 3 and not sr.passed:
                run.font.color.rgb = _rgb("C00000")

    doc.add_paragraph()

    overall_para = doc.add_paragraph(
        f"Overall validation status: "
        f"{'PASS — no flags raised.' if validation_report.overall_passed else f'REVIEW REQUIRED — {validation_report.total_flags} flag(s) raised across {len(validation_report.section_results)} sections.'}"
    )
    if overall_para.runs:
        overall_para.runs[0].bold = True
        overall_para.runs[0].font.size = Pt(STYLE.body_size)
        color = "007000" if validation_report.overall_passed else "C00000"
        overall_para.runs[0].font.color.rgb = _rgb(color)

    flagged_sections = [sr for sr in validation_report.section_results if sr.flagged]
    if flagged_sections:
        doc.add_paragraph()
        doc.add_heading("Flagged Values — Detail", level=2)
        for sr in flagged_sections:
            doc.add_paragraph(f"{sr.section_title}:", ).runs[0].bold = True
            for f in sr.flagged:
                raw_text = f.raw_text or "N/A"
                closest_kpi = f.closest_source_kpi or "unknown"
                closest_val = f.closest_source_value or "unknown"
                deviation = f"{f.deviation_pct:.1f}" if f.deviation_pct is not None else "N/A"
                doc.add_paragraph(
                    f"  • Value '{raw_text}' could not be matched to source data. "
                    f"Closest source: {closest_kpi} = {closest_val} "
                    f"(deviation: {deviation}%)",
                    style="List Bullet",
                )


def assemble_report(
    generated_report,       # GeneratedReport from pipeline.py
    validation_report,      # ValidationReport from validator.py
    section_plan,           # SectionPlan from section_planner.py
    output_path: Path | None = None,
) -> Path:
    """
    Build and save the Word document.

    Returns the Path to the saved .docx file.
    """
    company_name  = section_plan.company_name
    reporting_year = section_plan.reporting_year
    industry = section_plan.raw_data.get("industry_sector", "")

    # Output path
    if output_path is None:
        safe_name = re.sub(r"[^\w\-]", "_", company_name)
        filename = f"{safe_name}_{reporting_year}_sustainability_report.docx"
        output_path = settings.output_dir / filename

    output_path.parent.mkdir(parents=True, exist_ok=True)

    # Build document
    doc = Document()
    _set_margins(doc)
    _add_footer(doc, company_name)
    _add_title_page(doc, company_name, reporting_year, industry)
    _add_about_page(doc, company_name, reporting_year)

    # Add each section in template order
    section_text_map = generated_report.sections  

    for report_section in DOCUMENT_SECTION_ORDER:
        sid = report_section.section_id
        text = section_text_map.get(sid, "")
        if not text:
            continue
        _add_section(
            doc=doc,
            section_id=sid,
            title=report_section.title,
            gri_code=report_section.gri_code,
            tcfd_pillar=report_section.tcfd_pillar,
            body_text=text,
        )

    _add_validation_appendix(doc, validation_report)

    doc.save(str(output_path))
    return output_path